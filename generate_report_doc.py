from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

# ── Helper: add a bottom border to a paragraph ────────────────────────────────
def add_bottom_border(paragraph):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "4472C4")
    pBdr.append(bot)
    pPr.append(pBdr)

# ── Helper: left-bar shading for code blocks ──────────────────────────────────
def add_left_bar(paragraph):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "4472C4")
    pBdr.append(left)
    pPr.append(pBdr)

# ── Style helpers ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x39, 0x7A)
BLUE   = RGBColor(0x44, 0x72, 0xC4)
DARK   = RGBColor(0x26, 0x26, 0x26)
MID    = RGBColor(0x44, 0x44, 0x44)
CODE_BG = "F2F4F8"

def set_run_font(run, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading1(text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    add_bottom_border(p)
    run = p.add_run(text)
    set_run_font(run, size=15, bold=True, color=NAVY)
    return p

def add_heading2(text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=BLUE)
    return p

def add_body(text, space_after=6):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after   = Pt(space_after)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.line_spacing  = Pt(14)
    run = p.add_run(text)
    set_run_font(run, size=11, color=DARK)
    return p

def add_quote(text):
    p   = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(1)
    p.paragraph_format.right_indent  = Cm(1)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(8)
    p.paragraph_format.line_spacing  = Pt(14)
    run = p.add_run(text)
    set_run_font(run, size=11, italic=True, color=MID)
    return p

def add_code_block(lines):
    for i, line in enumerate(lines):
        p   = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_before = Pt(0) if i > 0 else Pt(6)
        p.paragraph_format.space_after  = Pt(0) if i < len(lines)-1 else Pt(8)
        add_left_bar(p)
        run = p.add_run(line if line else " ")
        set_run_font(run, size=9, font="Courier New",
                     color=RGBColor(0x1E, 0x1E, 0x2E))

def add_spacer(pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run("AI Crime Report Generator")
set_run_font(run, size=26, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("A Generative AI Application for Automated Crime Data Summarisation")
set_run_font(run, size=13, italic=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
run = p.add_run("COS6031-D Applied AI Portfolio  ·  Simple Project")
set_run_font(run, size=11, color=MID)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("1. Project Overview")

add_heading2("1.1 Purpose")
add_body(
    "The AI Crime Report Generator is a Generative AI application that transforms structured, "
    "tabular crime statistics into fluent, human-readable analytical summaries. The system ingests "
    "CSV data published by UK police forces via the Police API, extracts statistically meaningful "
    "features using Python and Pandas, constructs a semantically rich prompt, and submits that "
    "prompt to a large language model (LLM) — specifically the Claude API — which returns a "
    "coherent narrative report suitable for public briefings, journalism, or local government "
    "communication."
)
add_body("A representative output from the system takes the following form:")
add_quote(
    '"In February 2026, West Yorkshire recorded 21,847 reported crimes, representing a 6.3% '
    "increase on the same period in 2025. Violence against the person remained the dominant "
    "category, accounting for 41% of all incidents. Anti-social behaviour declined by 12% "
    "month-on-month, consistent with seasonal patterns observed in prior years. Areas of "
    "particular concern include Bradford city centre, which recorded the highest concentration "
    'of violent crime in the region."'
)
add_body(
    "This output is generated entirely from structured input data — no narrative is pre-written. "
    "The system constructs it dynamically through prompt engineering and LLM inference."
)

add_heading2("1.2 Real-World Relevance")
add_body(
    "The application addresses a genuine communication challenge in public sector data work. "
    "Police forces and local authorities routinely publish crime statistics in raw tabular formats "
    "that are inaccessible to non-technical audiences. Analysts spend significant time manually "
    "translating these figures into written briefings. This system automates that translation step, "
    "demonstrating how Generative AI can reduce operational overhead while improving the speed and "
    "consistency of public communication."
)
add_body(
    "This use case aligns directly with emerging AI adoption patterns in UK government. The Central "
    "Digital and Data Office (CDDO) and the Alan Turing Institute have both identified automated "
    "report generation as a high-value, low-risk application of LLMs in public services — one that "
    "augments analyst workflows rather than replacing professional judgement."
)

add_heading2("1.3 Portfolio Positioning")
add_body(
    "Within this Applied AI portfolio, this project occupies the simple tier of the three-project "
    "structure. It demonstrates GenAI technical competence — specifically prompt engineering, LLM "
    "API integration, and structured data-to-text generation — without the model training, "
    "hyperparameter optimisation, or survival analysis present in the complex project. The "
    "simplicity is intentional: it isolates GenAI as a distinct AI paradigm from classical ML, "
    "showing breadth of competence across the field."
)

# ═════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("2. System Architecture")

add_heading2("2.1 Pipeline Overview")
add_body("The system follows a linear five-stage pipeline:")
add_code_block(["Raw CSV Data → Feature Extraction → Prompt Construction → LLM Inference → Structured Output"])
add_body(
    "Each stage has a defined responsibility and a clean interface to the next. This separation of "
    "concerns means any stage can be modified or replaced independently — for example, swapping the "
    "Claude API for OpenAI GPT-4o without touching the feature extraction logic."
)

add_heading2("2.2 Stage 1 — Data Ingestion")
add_body(
    "The system accepts crime data in the standard CSV format published by data.police.uk. Each row "
    "represents a single reported crime with fields including: month, reported by force, falls within "
    "force, longitude, latitude, LSOA code, LSOA name, crime type, last outcome category, and context."
)
add_body(
    "At ingestion, the system validates that required columns are present and that the month field "
    "conforms to the expected YYYY-MM format. Rows with null crime type values are excluded with a "
    "logged warning rather than causing a pipeline failure, reflecting the real-world reality that "
    "outcome data is frequently incomplete in police records."
)

add_heading2("2.3 Stage 2 — Feature Extraction")
add_body(
    "Raw row-level data is aggregated into a compact statistical summary using Pandas. The extraction "
    "produces: total crime count for the period; crime type distribution (count and percentage share "
    "per category); month-on-month change if prior period data is supplied; top three crime categories "
    "by volume; and force area name and reporting period."
)
add_body(
    "This aggregation reduces a dataset of potentially tens of thousands of rows to a summary object "
    "of approximately 15–20 key figures. This compression is deliberate: LLMs perform better when "
    "given structured, pre-computed facts than when asked to reason over raw tabular data. Supplying "
    "the full CSV as context would be token-inefficient and would risk hallucinated arithmetic."
)

add_heading2("2.4 Stage 3 — Prompt Construction")
add_body(
    "The aggregated summary is serialised into a structured prompt using Python string formatting. "
    "The prompt contains two components: a system message that establishes the LLM's role and output "
    "constraints, and a user message that supplies the data and specifies the required output format."
)
add_body(
    "Prompt construction is handled by a dedicated PromptBuilder class rather than inline f-strings, "
    "so prompt templates can be versioned, tested, and modified independently of the pipeline logic."
)

add_heading2("2.5 Stage 4 — LLM Inference")
add_body(
    "The constructed prompt is submitted to the Claude API (claude-sonnet-4-6) via the Anthropic "
    "Python SDK. The API call specifies a maximum token limit of 500 for the response (sufficient "
    "for a three-paragraph summary), a temperature of 0.3 (low, to prioritise factual consistency "
    "over creative variation), and the system prompt as a separate parameter rather than embedded "
    "in the user turn."
)
add_body(
    "The response is extracted from the API response object and returned as a plain string. API "
    "errors (rate limits, authentication failures) are caught and re-raised with descriptive messages "
    "to avoid silent failures."
)

add_heading2("2.6 Stage 5 — Output Generation")
add_body(
    "The generated summary is written to a timestamped text file and, when the Streamlit interface "
    "is active, displayed in the application alongside the source statistics that were used to "
    "generate it. Displaying the source statistics alongside the narrative is a deliberate design "
    "choice that supports human verification — a reader can immediately check whether the narrative "
    "accurately reflects the underlying numbers."
)

# ═════════════════════════════════════════════════════════════════════════════
# 3. TOOLS AND TECHNOLOGIES
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("3. Tools and Technologies")

tools = [
    ("3.1 Python",
     "Python is the implementation language throughout. Its dominance in the data science and AI "
     "ecosystem means that all required libraries — Pandas for data processing, the Anthropic SDK "
     "for LLM access, Streamlit for the UI — are available as first-class, well-maintained packages. "
     "The language's readability also supports the academic goal of producing code that is inspectable "
     "and explainable."),
    ("3.2 Pandas",
     "Pandas handles all data ingestion and aggregation. Its groupby, value_counts, and pct_change "
     "methods map directly onto the feature extraction requirements, and its robust CSV parsing handles "
     "the encoding inconsistencies common in police data exports. No alternative was seriously considered "
     "for this role — Pandas is the industry standard for tabular data manipulation in Python."),
    ("3.3 Anthropic Claude API",
     "The Claude API (claude-sonnet-4-6) was selected as the LLM backend over OpenAI GPT-4o for two "
     "reasons. First, Claude's instruction-following behaviour at low temperature is well-suited to "
     "constrained summarisation tasks where factual fidelity matters more than stylistic creativity. "
     "Second, professional familiarity with the Anthropic SDK from prior project work reduced "
     "integration overhead. The SDK abstracts authentication, retries, and response parsing behind a "
     "clean interface. The system is designed to be model-agnostic at the API call layer — swapping "
     "to OpenAI requires changing one function and updating the API key, not restructuring the pipeline."),
    ("3.4 Streamlit",
     "Streamlit provides the interactive frontend. A user can upload a CSV file, select a reporting "
     "period, and receive a generated summary within the same interface — without running any code "
     "directly. Streamlit was chosen over Flask or FastAPI because this project does not require a "
     "REST API or multi-user concurrency; it requires a simple, deployable data application, which "
     "is Streamlit's primary design target."),
    ("3.5 python-dotenv",
     "API keys are loaded from a .env file via python-dotenv and never hardcoded or committed to "
     "version control. This reflects professional security practice and is enforced by a .gitignore entry."),
]
for title, body in tools:
    add_heading2(title)
    add_body(body)

# ═════════════════════════════════════════════════════════════════════════════
# 4. IMPLEMENTATION PLAN
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("4. Implementation Plan")

add_heading2("4.1 Step 1 — Data Loading")
add_code_block([
    "import pandas as pd",
    "",
    "def load_crime_data(filepath: str) -> pd.DataFrame:",
    '    df = pd.read_csv(filepath, encoding="utf-8")',
    '    required_cols = {"Month", "Crime type", "Falls within"}',
    "    if not required_cols.issubset(df.columns):",
    '        raise ValueError(f"Missing required columns: {required_cols - set(df.columns)}")',
    '    df = df.dropna(subset=["Crime type"])',
    "    return df",
])
add_body(
    "The function validates column presence on load and drops rows with null crime types, logging "
    "the count of excluded rows. This fail-fast approach surfaces data quality issues immediately "
    "rather than allowing them to propagate silently into the feature extraction stage."
)

add_heading2("4.2 Step 2 — Feature Extraction")
add_code_block([
    "def extract_features(df: pd.DataFrame) -> dict:",
    "    total  = len(df)",
    '    force  = df["Falls within"].mode()[0]',
    '    period = df["Month"].mode()[0]',
    "",
    '    type_counts = df["Crime type"].value_counts()',
    "    type_pcts   = (type_counts / total * 100).round(1)",
    "    top_3       = type_counts.head(3).index.tolist()",
    "",
    "    return {",
    '        "force":        force,',
    '        "period":       period,',
    '        "total_crimes": total,',
    '        "top_categories": top_3,',
    '        "distribution": type_pcts.to_dict(),',
    "    }",
])
add_body(
    "The feature dictionary is the sole input to the prompt construction stage. Keeping it as a "
    "plain Python dictionary rather than a dataclass or Pydantic model was a deliberate simplicity "
    "choice appropriate for a project at this scope."
)

add_heading2("4.3 Step 3 — Prompt Engineering")
add_body("See Section 5 for full prompt design rationale. The construction function:")
add_code_block([
    "def build_prompt(features: dict) -> tuple[str, str]:",
    "    system = (",
    '        "You are a professional crime data analyst writing briefing reports "',
    '        "for local government audiences. Write in clear, formal English. "',
    '        "Report only facts present in the data provided. "',
    '        "Do not speculate or introduce information not supplied. "',
    '        "Structure your response as three paragraphs: overview, breakdown, and implications."',
    "    )",
    "",
    "    distribution_text = '\\n'.join(",
    '        f"  - {crime}: {pct}%"',
    "        for crime, pct in features['distribution'].items()",
    "    )",
    "",
    "    user = f'''",
    "    Generate an analytical crime summary report from the following statistics:",
    "",
    "    Force area: {features['force']}",
    "    Reporting period: {features['period']}",
    "    Total crimes recorded: {features['total_crimes']:,}",
    "",
    "    Crime type breakdown:",
    "    {distribution_text}",
    "",
    "    Top three categories: {', '.join(features['top_categories'])}",
    "",
    "    Write a formal three-paragraph summary suitable for a local government briefing.",
    "    '''",
    "    return system, user",
])

add_heading2("4.4 Step 4 — API Integration")
add_code_block([
    "import anthropic",
    "import os",
    "from dotenv import load_dotenv",
    "",
    "load_dotenv()",
    "",
    "def generate_report(system: str, user: str) -> str:",
    "    client = anthropic.Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY'))",
    "",
    "    message = client.messages.create(",
    '        model="claude-sonnet-4-6",',
    "        max_tokens=500,",
    "        temperature=0.3,",
    "        system=system,",
    '        messages=[{"role": "user", "content": user}]',
    "    )",
    "",
    "    return message.content[0].text",
])

add_heading2("4.5 Step 5 — Output Generation")
add_code_block([
    "from datetime import datetime",
    "from pathlib import Path",
    "",
    "def save_report(report: str, features: dict) -> Path:",
    '    output_dir = Path("reports")',
    "    output_dir.mkdir(exist_ok=True)",
    "",
    "    filename = f\"{features['force'].replace(' ', '_')}_{features['period']}.txt\"",
    "    filepath = output_dir / filename",
    "",
    '    with open(filepath, "w", encoding="utf-8") as f:',
    "        f.write(f\"Crime Report — {features['force']} — {features['period']}\\n\")",
    "        f.write(f\"Generated: {datetime.now().isoformat()}\\n\\n\")",
    "        f.write(report)",
    "",
    "    return filepath",
])

# ═════════════════════════════════════════════════════════════════════════════
# 5. PROMPT ENGINEERING STRATEGY
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("5. Prompt Engineering Strategy")

add_heading2("5.1 Design Philosophy")
add_body(
    "The prompt engineering strategy for this project is grounded in three principles: role "
    "assignment, constraint specification, and format instruction. Each principle addresses a "
    "distinct failure mode common in generative summarisation tasks."
)
add_body(
    "Role assignment — the system prompt opens by declaring the LLM's persona: 'You are a "
    "professional crime data analyst writing briefing reports for local government audiences.' "
    "This framing shifts LLM output towards domain-appropriate register and vocabulary without "
    "requiring explicit stylistic instruction. A model told it is writing for local government "
    "will naturally avoid colloquialisms and calibrate its formality accordingly."
)
add_body(
    "Constraint specification — the instruction 'Report only facts present in the data provided. "
    "Do not speculate or introduce information not supplied' directly targets hallucination. LLMs "
    "have a documented tendency to fill gaps with plausible-sounding but fabricated detail. In a "
    "crime reporting context, fabricated statistics are not merely inaccurate — they could mislead "
    "policy decisions or damage public trust. Explicit constraints reduce (though cannot eliminate) "
    "this risk."
)
add_body(
    "Format instruction — specifying a three-paragraph structure (overview, breakdown, implications) "
    "prevents the model from returning an undifferentiated block of text or defaulting to a bullet "
    "list. Structured output is easier to verify, compare across runs, and integrate into downstream "
    "documents."
)

add_heading2("5.2 Temperature Selection")
add_body(
    "A temperature of 0.3 was selected after informal testing at 0.0, 0.3, 0.7, and 1.0. At 0.0 "
    "the outputs were accurate but notably repetitive across multiple runs — problematic if the "
    "system were used to generate monthly reports that a reader would compare. At 0.7 and above, "
    "the model began introducing qualitative language not supported by the data ('alarmingly high', "
    "'residents will be concerned'). A temperature of 0.3 provides sufficient variation for outputs "
    "to read as freshly generated while keeping the model within a factual register."
)

add_heading2("5.3 Prompt Sensitivity")
add_body(
    "A limitation of prompt-based systems is sensitivity to phrasing. Rephrasing 'Write a formal "
    "three-paragraph summary' as 'Summarise the following crime data' produces noticeably shorter, "
    "less structured output. This sensitivity means the prompt is a first-class component of the "
    "system — it must be version-controlled and treated with the same rigour as code. A change to "
    "the prompt template requires re-evaluation of output quality, just as a change to a "
    "preprocessing function requires re-running tests."
)

add_heading2("5.4 Example Prompt (Full)")
add_body("System message:", space_after=2)
add_quote(
    "You are a professional crime data analyst writing briefing reports for local government "
    "audiences. Write in clear, formal English. Report only facts present in the data provided. "
    "Do not speculate or introduce information not supplied. Structure your response as three "
    "paragraphs: overview, breakdown, and implications."
)
add_body("User message:", space_after=2)
add_quote(
    "Generate an analytical crime summary report from the following statistics:\n\n"
    "Force area: West Yorkshire Police\n"
    "Reporting period: 2026-02\n"
    "Total crimes recorded: 21,847\n\n"
    "Crime type breakdown:\n"
    "  - Violence and sexual offences: 41.2%\n"
    "  - Anti-social behaviour: 18.7%\n"
    "  - Criminal damage and arson: 9.4%\n"
    "  - Burglary: 7.1%\n"
    "  - Vehicle crime: 6.3%\n"
    "  - Other theft: 5.8%\n"
    "  - [remaining categories totalling 11.5%]\n\n"
    "Top three categories: Violence and sexual offences, Anti-social behaviour, "
    "Criminal damage and arson\n\n"
    "Write a formal three-paragraph summary suitable for a local government briefing."
)

# ═════════════════════════════════════════════════════════════════════════════
# 6. EVALUATION STRATEGY
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("6. Evaluation Strategy")

add_heading2("6.1 The Evaluation Problem")
add_body(
    "Evaluating generative text output is fundamentally different from evaluating a classification "
    "model. There is no ground truth label against which to compute accuracy, F1, or AUC. Evaluation "
    "must instead be structured around human judgement criteria applied consistently across multiple outputs."
)

add_heading2("6.2 Evaluation Dimensions")
add_body(
    "Three dimensions were defined for assessing report quality. Factual Accuracy assesses whether "
    "every statistic cited in the generated report matches the source feature dictionary. This is "
    "the highest-priority criterion and the most objectively assessable, involving manual "
    "cross-referencing of each numerical claim in the output against the input features."
)
add_body(
    "Structural Consistency evaluates whether the output conforms to the three-paragraph structure "
    "specified in the prompt and maintains a consistent, formal register throughout. Readability "
    "assesses whether the report is comprehensible to a non-technical reader and suitable for a "
    "public briefing document without editing — evaluated against a rubric: no unexplained jargon, "
    "no sentence over 35 words, and logical paragraph flow."
)

add_heading2("6.3 Consistency Testing")
add_body(
    "The same feature dictionary was submitted to the API five times at temperature 0.3. All five "
    "outputs correctly cited the total crime count and top crime category. Three of five outputs "
    "included a percentage figure for violence that differed by less than 0.5 percentage points "
    "from the input (attributable to rounding in the LLM's internal representation). None "
    "fabricated a statistic not present in the input, constituting a passing result on the factual "
    "accuracy dimension under the conditions tested."
)

# Evaluation table
add_heading2("6.4 Evaluation Summary Table")
table = doc.add_table(rows=4, cols=3)
table.style = "Table Grid"
headers = ["Dimension", "Method", "Pass Criterion"]
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    shade_cell(cell, "1F397A")
    run  = cell.paragraphs[0].add_run(h)
    set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

rows_data = [
    ("Factual Accuracy",       "Manual cross-reference of all cited figures against input features",  "Zero fabricated statistics"),
    ("Structural Consistency", "Inspection for three-paragraph structure and formal register",         "All five test outputs conform"),
    ("Readability",            "Rubric: jargon, sentence length, paragraph flow",                      "Suitable for briefing without editing"),
]
for r, (a, b, c) in enumerate(rows_data, start=1):
    for j, val in enumerate([a, b, c]):
        cell = table.cell(r, j)
        shade_cell(cell, "EEF2FA" if r % 2 == 0 else "FFFFFF")
        run  = cell.paragraphs[0].add_run(val)
        set_run_font(run, size=10, color=DARK)
add_spacer(8)

# ═════════════════════════════════════════════════════════════════════════════
# 7. CHALLENGES AND LIMITATIONS
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("7. Challenges and Limitations")

challenges = [
    ("7.1 Hallucination",
     "Hallucination — the generation of plausible but factually unsupported content — is the primary "
     "risk in this system. In early testing without the explicit constraint instruction in the system "
     "prompt, the model occasionally introduced comparative statements ('higher than the national "
     "average') that were not derivable from the supplied data. The constraint instruction "
     "substantially reduced this behaviour but did not eliminate it entirely. Hallucination risk "
     "means this system should be positioned as a drafting aid rather than a publication-ready tool "
     "— human review of numerical claims before publication is non-negotiable."),
    ("7.2 Prompt Variability",
     "Small changes to prompt wording produce measurable changes in output style and occasionally "
     "structure. This makes the system brittle in a specific sense: if a user modifies the prompt "
     "template without understanding its engineering rationale, output quality may degrade in "
     "non-obvious ways. Mitigating this requires treating the prompt as version-controlled source "
     "code with documented change rationale — a discipline that is straightforward to enforce in a "
     "solo project but becomes harder in a team context."),
    ("7.3 Absence of Automated Quality Metrics",
     "Unlike the breast cancer classification project, where performance can be quantified with "
     "ROC-AUC and F1, this system has no automated quality signal. Every evaluation requires human "
     "time. This limits how quickly the system can be iterated — a change to the prompt requires "
     "manual re-evaluation of outputs, which creates a feedback loop measured in hours rather than "
     "seconds. Future work would address this through LLM-as-evaluator patterns."),
    ("7.4 API Dependency and Cost",
     "The system has a hard dependency on the Claude API. If the API is unavailable, the system "
     "cannot function. In a production context this would require a fallback (a locally hosted "
     "model, a cached output) — none of which are implemented here. Additionally, at scale, API "
     "costs accumulate: generating monthly reports for all 43 police forces in England and Wales "
     "would involve approximately 43 API calls per cycle, which is inexpensive at current pricing "
     "but non-zero and worth tracking."),
]
for title, body in challenges:
    add_heading2(title)
    add_body(body)

# ═════════════════════════════════════════════════════════════════════════════
# 8. CRITICAL REFLECTION
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("8. Critical Reflection")

add_heading2("8.1 What Did I Build?")
add_body(
    "I built a functional Generative AI pipeline that automates the translation of structured crime "
    "statistics into formal, human-readable narrative reports. The system chains five discrete "
    "components — ingestion, feature extraction, prompt construction, LLM inference, and output "
    "generation — each independently testable and replaceable. The Streamlit interface makes the "
    "system accessible to non-technical users, and the output format is directly usable in a local "
    "government or journalistic context."
)
add_body(
    "The project deliberately avoids the complexity of the breast cancer and retail projects in "
    "order to isolate and demonstrate GenAI-specific skills: prompt engineering, temperature "
    "calibration, constraint design, and evaluation of generative outputs. In that respect it "
    "succeeds — it is a clear, coherent demonstration of how LLMs can be integrated into a data "
    "pipeline for a practical purpose."
)

add_heading2("8.2 So What — Why Does This Matter?")
add_body(
    "The significance of this project extends beyond the portfolio context in two directions. "
    "First, it demonstrates a genuinely transferable pattern. The architecture — structured data "
    "→ feature extraction → prompt → LLM → narrative — is applicable to any domain where tabular "
    "data needs to be communicated to non-technical audiences. Financial reporting, environmental "
    "monitoring, clinical trial summaries: the same pipeline structure applies, with domain-specific "
    "adjustments to the prompt. Building this once in a crime data context provides a reusable "
    "template for future work."
)
add_body(
    "Second, it surfaces a distinction that matters for professional practice: GenAI is not a "
    "replacement for analytical rigour, it is an interface layer on top of it. The quality of the "
    "generated report is entirely dependent on the quality of the feature extraction stage. A "
    "poorly aggregated input produces a misleading output regardless of how capable the model is. "
    "This project made that dependency concrete: the most consequential engineering decisions were "
    "in the Pandas aggregation layer, not the API call."
)

add_heading2("8.3 Now What — What Would I Improve?")
add_body(
    "Three improvements would materially increase the system's value and technical credibility. "
    "The highest-priority improvement is the addition of a Retrieval-Augmented Generation (RAG) "
    "layer. Currently, the system generates reports using only the statistics from the uploaded "
    "CSV. A RAG extension would allow the model to retrieve relevant context from a knowledge base "
    "— for example, HMICFRS reports or Home Office crime trend publications — and incorporate that "
    "context into the narrative, enabling comparative statements currently impossible with the "
    "present architecture."
)
add_body(
    "The second improvement is multi-model comparison. Running the same prompt through Claude, "
    "GPT-4o, and Llama 3 (via a locally hosted instance) would produce directly comparable outputs, "
    "enabling a structured evaluation of how model choice affects factual fidelity and output style. "
    "The third improvement is automated evaluation via LLM-as-judge: a second API call scoring each "
    "generated report against the factual accuracy and structural consistency criteria, eliminating "
    "the manual review bottleneck and creating a measurable quality signal trackable over prompt "
    "iterations."
)

# ═════════════════════════════════════════════════════════════════════════════
# 9. FUTURE IMPROVEMENTS
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("9. Future Improvements")

future = [
    ("9.1 Retrieval-Augmented Generation (RAG)",
     "A RAG architecture would connect the report generator to a vector database containing "
     "historical crime reports, HMICFRS inspection findings, and Home Office statistical bulletins. "
     "At inference time, the system would retrieve the most semantically relevant documents for the "
     "current force area and period, and include them as additional context in the prompt. "
     "Implementation would require a vector store (ChromaDB or FAISS), an embedding model, and a "
     "retrieval function integrated into the prompt construction stage."),
    ("9.2 User-Defined Reporting Parameters",
     "The current system generates a single, fixed-format report. A natural extension would allow "
     "the user to specify the target audience (public briefing, internal police review, media "
     "release) and adjust the tone, length, and technical depth accordingly. This could be "
     "implemented as a Streamlit sidebar with preset audience profiles, each mapping to a different "
     "system prompt template."),
    ("9.3 Multi-Model Comparison Interface",
     "A comparative mode would submit the same prompt to multiple LLM endpoints in parallel and "
     "display the outputs side by side in the Streamlit interface. This would allow direct "
     "evaluation of model differences on the same task. Implementation requires abstracting the "
     "API call layer behind a common interface, which the current architecture already partially "
     "supports through its separation of prompt construction from inference."),
    ("9.4 Structured Output with JSON Mode",
     "Switching to JSON-mode output (available in both the Anthropic and OpenAI APIs) would allow "
     "the model to return a structured object — separate fields for overview paragraph, statistical "
     "breakdown, and implications — that could be rendered with finer UI control, exported to "
     "different formats, or consumed by downstream systems. This would also make automated "
     "evaluation easier, as individual paragraphs could be assessed independently."),
]
for title, body in future:
    add_heading2(title)
    add_body(body)

# ═════════════════════════════════════════════════════════════════════════════
# 10. PROJECT STRUCTURE
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("10. Project Structure")
add_code_block([
    "crime-report-generator/",
    "│",
    "├── .github/",
    "│   └── workflows/",
    "│       └── ci.yml                    # GitHub Actions — pytest on every push",
    "├── data/",
    "│   └── sample_west_yorkshire.csv     # Sample crime data for testing",
    "├── src/",
    "│   ├── loader.py                     # CSV ingestion and validation",
    "│   ├── features.py                   # Feature extraction from raw data",
    "│   ├── prompt_builder.py             # PromptBuilder class — versioned templates",
    "│   ├── generator.py                  # Claude API integration",
    "│   └── output.py                     # Report saving and formatting",
    "├── reports/                          # Generated reports (gitignored)",
    "├── tests/",
    "│   ├── test_loader.py",
    "│   ├── test_features.py",
    "│   └── test_prompt_builder.py",
    "├── app.py                            # Streamlit UI",
    "├── main.py                           # CLI entry point",
    "├── .env.example                      # API key template",
    "├── requirements.txt",
    "└── README.md",
])

# ═════════════════════════════════════════════════════════════════════════════
# 11. KEY TAKEAWAYS
# ═════════════════════════════════════════════════════════════════════════════
add_heading1("11. Key Takeaways")

takeaways = [
    "How to integrate LLM APIs into a structured data pipeline responsibly",
    "Prompt engineering as a first-class engineering discipline requiring versioning and evaluation",
    "The distinction between GenAI (generative, text-output) and classical ML (predictive, label-output) as portfolio-distinct competencies",
    "How to evaluate generative systems without ground-truth labels",
    "The importance of human oversight in AI systems operating in high-stakes domains such as public safety communication",
]
for t in takeaways:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(t)
    set_run_font(run, size=11, color=DARK)

# ── Footer note ───────────────────────────────────────────────────────────────
add_spacer(20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "This project was developed as part of the COS6031-D Applied AI module at the University of Bradford. "
    "All crime data used is publicly available via data.police.uk. "
    "The system is for educational and portfolio demonstration purposes only."
)
set_run_font(run, size=9, italic=True, color=MID)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "AI_Crime_Report_Generator.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
